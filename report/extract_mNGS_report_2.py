#!/usr/bin/env python3
# -*- coding: utf-8 -*
# @Author  : yellow_huang

import os,glob
import pandas as pd
from path import  Path
from docx import Document
import re


in_path = '/mnt/home/huanggy/project/report/LX2110650-吴伟双.docx'
# 按二维矩阵索引遍历表格
doc = Document(in_path)
tables = doc.tables
print(tables)

DNA_tb_num=3 #DNA 病毒
bac_tb_num=4
fungi_tb_num=5
pro_tb_num=6
tolarent_gene_tb_num=7
likely_bacground_tb_num=8
mytable = tables[DNA_tb_num]


# print(f'第一个表格: {len(mytable.rows)} 行 X {len(mytable.columns)}列')
# for i, row in enumerate(mytable.rows):
#     print(f'第 {i+1} 行有 {len(row.cells)} 列')
#     tmp = []
#     for j, cell in enumerate(row.cells):
#         tmp.append(cell.text)
#
#     if len(tmp) >=7:
#         print("|".join(tmp))
#     else:
#         tmp.insert(0,"类型")
#         print("|".join(tmp))
#         # if cell not in cell_set:
#         #     cell_set.add(cell)
#         #     cell.text += 'test'
#         #     print('=====>',cell.text)


def run(dir_path,out):
    file_list = glob.glob(Path(dir_path)/"*.docx")
    print(file_list)
    outfile = open(out, 'w')
    print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%",sep="\t",file=outfile)
    for f in file_list:
        doc = Document(f)
        file_name = os.path.basename(f).replace(".docx","")
        tables = doc.tables

        table_num = [3,4,5,6,8]
        for num in table_num:
            mytable = tables[num]

            for i, row in enumerate(mytable.rows):
                #print(f'第 {i+1} 行有 {len(row.cells)} 列')
                #第一行，属+终
                #第二行：类型	名称	序列数	相对丰度%	名称	序列数	相对丰度%
                if i == 0 :
                    continue
                if mytable.cell(i, 0).text == "名称" or mytable.cell(i, 0).text == "类型":
                    continue

                ##内容——
                # if mytable.cell(i, 0).text == "--":
                #     continue
                tmp = []
                for j, cell in enumerate(row.cells):
                    tmp.append(cell.text)
                if len(tmp) >= 7:
                    #print(file_name, "\t".join(tmp))
                    if len(set(tmp)) == 1 and list(set(tmp))[0] == "--":
                        continue
                    print(file_name,"\t".join(tmp).replace("\n","_"),sep="\t",file=outfile)
                else:
                    if len(set(tmp)) == 1 and list(set(tmp))[0] == "--":
                        continue
                    tmp.insert(0, "-")
                    #print(file_name, "\t".join(tmp))
                    print(file_name,"\t".join(tmp).replace("\n","_"),sep="\t",file=outfile)


    outfile.close()


#run(dir_path="/mnt/home/huanggy/project/report",out='/mnt/home/huanggy/project/report_out.txt')

#20211020   compare report 与kraken
def yanzhen_species(file_path,quire_name):
    with open(file_path, 'r') as fh:
        next(fh)
        for line in fh:
            if not line:
                break
            else:
                line = line.strip()
                if re.search(quire_name,line):
                    name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, added_reads, new_est_reads, fraction_total_reads = line.strip().split("\t")
                    print("======>", name,taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads)
                    return name,taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads
                else:
                    pass


def yanzhen_species_blast(file_path,quire_name):
    with open(file_path, 'r') as fh:
        next(fh)
        for line in fh:
            if not line:
                break
            else:
                line = line.strip()
                if re.search(quire_name,line):
                    print(file_path)
                    taxid, lineage_id, lineage_name, count, s_abundance = line.strip().split(",")
                    #print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, name,taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads,file=out1, sep='\t')
                    return taxid, lineage_id, lineage_name, count, s_abundance
                else:
                    pass

"""
report_tb =pd.read_table("/mnt/home/huanggy/project/report_out.txt",sep="\t",header=0)

out1 = open("/mnt/home/huanggy/project/report_out_compaire_s_2.txt",'w')
out2 = open("/mnt/home/huanggy/project/report_out_compaire_g_2.txt",'w')
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads",sep="\t",file=out1)
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads",sep="\t",file=out2)
for ind ,row in report_tb.iterrows():
    id = row["id"]
    gelanshi = row['类型']
    genus_name = row["属名"].split("_")[1] if len(row["属名"].split("_")) >=2 else row["属名"].split("_")[0]
    genus_cnt = row['属序列数']
    genus_abudance = row['属相对丰度%']
    s_name = row['种名'].split("_")[1] if len(row['种名'].split("_")) >=2 else row['种名'].split("_")[0]

    s_cnt = row['种序列数']
    s_abudance = row['种相对丰度%']

    sid = id.split("-")[0]


    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test_compare/result/kraken2")/f"{sid}.s.bracken"):
        #print(Path("/mnt/home/huanggy/project/20210927_test_compare/result/kraken2")/f"{sid}.s.bracken",s_name)
        res = yanzhen_species(file_path=Path("/mnt/home/huanggy/project/20210927_test_compare/result/kraken2")/f"{sid}.s.bracken",quire_name=s_name)
        if res:
            print("species True",res)
            name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads = res
            print(name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads)
        else:
            print(res)
            name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads = "-","-", "-", "-", "-", "-"
        print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, name, taxonomy_id,taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads, file=out1, sep='\t')


    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test_compare/result/kraken2")/f"{sid}.g.bracken"):
        res = yanzhen_species(file_path=Path("/mnt/home/huanggy/project/20210927_test_compare/result/kraken2") / f"{sid}.g.bracken",quire_name=genus_name)
        if res:
            print("genus True")
            name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads = res
        else:
            name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads = "-","-", "-", "-", "-", "-"
        print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, name, taxonomy_id,taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads, file=out2, sep='\t')

out1.close()
out2.close()
"""


#20211021 compare report 与kraken 与blast

tb1 = pd.read_table("/mnt/home/huanggy/project/report_out_compaire_s_2.txt",sep="\t",header=0)
out1 = open("/mnt/home/huanggy/project/compare_report_kraken2_blast.s_2.txt",'w')
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads","lineage_id","lineage_name","count","s_abundance",sep="\t",file=out1)
for ind ,row in tb1.iterrows():
    id = row["id"]
    gelanshi = row['类型']
    genus_name = row["属名"]
    genus_cnt = row['属序列数']
    genus_abudance = row['属相对丰度%']
    s_name = row['种名']

    s_cnt = row['种序列数']
    s_abudance = row['种相对丰度%']
    kraken_classify_name = row['kraken_classify_name']
    taxonomy_id = row['taxonomy_id']
    taxonomy_lvl = row['taxonomy_lvl']
    kraken_assigned_reads = row['kraken_assigned_reads']
    new_est_reads = row['new_est_reads']
    fraction_total_reads = row['fraction_total_reads']
    sid = id.split("-")[0]

    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test_compare/result/align")/f"{sid}.s.csv"):
        res = yanzhen_species_blast(file_path=Path("/mnt/home/huanggy/project/20210927_test_compare/result/align")/f"{sid}.s.csv",quire_name=s_name)
        if res:
            taxid, lineage_id, lineage_name, count, s_abundance = res
        else:
            taxid, lineage_id, lineage_name, count, s_abundance = "-","-", "-", "-", "-"

        print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, kraken_classify_name, taxonomy_id,
              taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads, lineage_id, lineage_name, count, s_abundance,file=out1, sep='\t')
out1.close()

tb2 = pd.read_table("/mnt/home/huanggy/project/report_out_compaire_g_2.txt",sep="\t",header=0)
out2 = open("/mnt/home/huanggy/project/compare_report_kraken2_blast.g_2.txt",'w')
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads","lineage_id","lineage_name","count","g_abundance",sep="\t",file=out2)

for ind ,row in tb1.iterrows():
    id = row["id"]
    gelanshi = row['类型']
    genus_name = row["属名"]
    genus_cnt = row['属序列数']
    genus_abudance = row['属相对丰度%']
    s_name = row['种名']

    s_cnt = row['种序列数']
    s_abudance = row['种相对丰度%']
    kraken_classify_name = row['kraken_classify_name']
    taxonomy_id = row['taxonomy_id']
    taxonomy_lvl = row['taxonomy_lvl']
    kraken_assigned_reads = row['kraken_assigned_reads']
    new_est_reads = row['new_est_reads']
    fraction_total_reads = row['fraction_total_reads']
    sid = id.split("-")[0]

    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test_compare/result/align")/f"{sid}.s.csv"):
        res = yanzhen_species_blast(file_path=Path("/mnt/home/huanggy/project/20210927_test_compare/result/align")/f"{sid}.g.csv",quire_name=genus_name)
        if res:
            taxid, lineage_id, lineage_name, count, g_abundance = res
        else:
            taxid, lineage_id, lineage_name, count, g_abundance = "-","-", "-", "-", "-"

        print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, kraken_classify_name, taxonomy_id,
              taxonomy_lvl, kraken_assigned_reads, new_est_reads, fraction_total_reads, lineage_id, lineage_name, count, g_abundance,file=out2, sep='\t')
out1.close()
out2.close()





