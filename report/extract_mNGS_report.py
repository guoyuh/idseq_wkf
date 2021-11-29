#!/usr/bin/env python3
# -*- coding: utf-8 -*
# @Author  : yellow_huang

import os,glob
import pandas as pd
from path import  Path
from docx import Document



in_path = '/mnt/home/huanggy/project/report/LX2110650-吴伟双.docx'
# 按二维矩阵索引遍历表格
doc = Document(in_path)
tables = doc.tables
print(tables)

DNA_tb_num=3 #DNA 病毒
bac_tb_num=4
fungi_tb_num=5
pro_tb_num=6
tolarent_gene_tb_num=7
likely_bacground_tb_num=8
mytable = tables[DNA_tb_num]


# print(f'第一个表格: {len(mytable.rows)} 行 X {len(mytable.columns)}列')
# for i, row in enumerate(mytable.rows):
#     print(f'第 {i+1} 行有 {len(row.cells)} 列')
#     tmp = []
#     for j, cell in enumerate(row.cells):
#         tmp.append(cell.text)
#
#     if len(tmp) >=7:
#         print("|".join(tmp))
#     else:
#         tmp.insert(0,"类型")
#         print("|".join(tmp))
#         # if cell not in cell_set:
#         #     cell_set.add(cell)
#         #     cell.text += 'test'
#         #     print('=====>',cell.text)


def run(dir_path,out):
    file_list = glob.glob(Path(dir_path)/"*.docx")
    print(file_list)
    outfile = open(out, 'w')
    print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%",sep="\t",file=outfile)
    for f in file_list:
        doc = Document(f)
        file_name = os.path.basename(f).replace(".docx","")
        tables = doc.tables

        table_num = [3,4,5,6,8]
        for num in table_num:
            mytable = tables[num]

            for i, row in enumerate(mytable.rows):
                #print(f'第 {i+1} 行有 {len(row.cells)} 列')
                #第一行，属+终
                #第二行：类型	名称	序列数	相对丰度%	名称	序列数	相对丰度%
                if i == 0 :
                    continue
                if mytable.cell(i, 0).text == "名称" or mytable.cell(i, 0).text == "类型":
                    continue

                ##内容——
                # if mytable.cell(i, 0).text == "--":
                #     continue
                tmp = []
                for j, cell in enumerate(row.cells):
                    tmp.append(cell.text)
                if len(tmp) >= 7:
                    #print(file_name, "\t".join(tmp))
                    if len(set(tmp)) == 1 and list(set(tmp))[0] == "--":
                        continue
                    print(file_name,"\t".join(tmp).replace("\n",""),sep="\t",file=outfile)
                else:
                    if len(set(tmp)) == 1 and list(set(tmp))[0] == "--":
                        continue
                    tmp.insert(0, "-")
                    #print(file_name, "\t".join(tmp))
                    print(file_name,"\t".join(tmp).replace("\n",""),sep="\t",file=outfile)


    outfile.close()


#run(dir_path="/mnt/home/huanggy/project/report",out='/mnt/home/huanggy/project/report_out.txt')

report_tb =pd.read_table("/mnt/home/huanggy/project/report_out.txt",sep="\t",header=0)


out1 = open("/mnt/home/huanggy/project/report_out_compaire_s.txt",'w')
out2 = open("/mnt/home/huanggy/project/report_out_compaire_g.txt",'w')
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads",sep="\t",file=out1)
print("id","类型","属名","属序列数","属相对丰度%","种名",'种序列数',"种相对丰度%","kraken_classify_name","taxonomy_id","taxonomy_lvl" ,"kraken_assigned_reads","new_est_reads","fraction_total_reads",sep="\t",file=out2)
for ind ,row in report_tb.iterrows():
    id = row["id"]
    gelanshi = row['类型']
    genus_name = row["属名"]
    genus_cnt = row['属序列数']
    genus_abudance = row['属相对丰度%']
    s_name = row['种名']
    s_cnt = row['种序列数']
    s_abudance = row['种相对丰度%']

    sid = id.split("-")[0]


    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test/result/kraken2")/f"{sid}.s.bracken"):
        with open(Path("/mnt/home/huanggy/project/20210927_test/result/kraken2")/f"{sid}.s.bracken",'r') as fh:
            next(fh)
            for line in fh:
                if not line :
                    break
                else:
                    name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, added_reads, new_est_reads, fraction_total_reads = line.split("\t")
                    if name in s_name:
                        print(id,gelanshi,genus_name,genus_cnt,genus_abudance,s_name,s_cnt,s_abudance,name,taxonomy_id,taxonomy_lvl,kraken_assigned_reads,new_est_reads,fraction_total_reads,file=out1,sep='\t')
                        break
                    else:
                        #pass
                        print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, "-","-", "-", "-", "-", "-",file=out1, sep='\t')

    if os.path.exists(Path("/mnt/home/huanggy/project/20210927_test/result/kraken2")/f"{sid}.g.bracken"):
        with open(Path("/mnt/home/huanggy/project/20210927_test/result/kraken2")/f"{sid}.g.bracken",'r') as fh:
            next(fh)
            for line in fh:
                name, taxonomy_id, taxonomy_lvl, kraken_assigned_reads, added_reads, new_est_reads, fraction_total_reads = line.split("\t")
                if name in genus_name:
                    print(id,gelanshi,genus_name,genus_cnt,genus_abudance,s_name,s_cnt,s_abudance,name,taxonomy_id,taxonomy_lvl,kraken_assigned_reads,new_est_reads,fraction_total_reads,file=out2,sep='\t')
                else:
                    #pass
                    print(id, gelanshi, genus_name, genus_cnt, genus_abudance, s_name, s_cnt, s_abudance, "-","-", "-", "-", "-", "-",file=out1, sep='\t')


out1.close()
out2.close()




